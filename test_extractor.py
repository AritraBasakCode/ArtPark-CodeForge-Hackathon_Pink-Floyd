"""
tests/test_extractor.py
-----------------------
Run with:  python -m pytest tests/ -v
"""

import os
import sys
import pytest
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
from extractor import (
    extract_resume_text,
    _clean_text,
    _split_into_sections,
    _detect_format,
    EmptyDocumentError,
    UnsupportedFormatError,
    ResumeExtractionError,
)


# ---------------------------------------------------------------------------
# _clean_text
# ---------------------------------------------------------------------------

class TestCleanText:

    def test_removes_bullet_unicode(self):
        raw = "• Python\n• JavaScript\n\u2022 Java"
        result = _clean_text(raw)
        assert "•" not in result
        assert "\u2022" not in result
        assert "Python" in result

    def test_normalizes_dashes(self):
        raw = "Senior Engineer\u2014TechCorp\nJan 2020\u2013Dec 2022"
        result = _clean_text(raw)
        assert "\u2014" not in result
        assert "\u2013" not in result

    def test_collapses_excess_blank_lines(self):
        raw = "Section A\n\n\n\n\nSection B"
        result = _clean_text(raw)
        # Should not have more than 2 consecutive blank lines (3 newlines = 2 blanks)
        assert "\n\n\n\n" not in result

    def test_strips_trailing_whitespace(self):
        raw = "Python   \nJavaScript   \n"
        result = _clean_text(raw)
        for line in result.splitlines():
            assert line == line.rstrip()

    def test_handles_non_breaking_space(self):
        raw = "React\xa0Native"
        result = _clean_text(raw)
        assert "\xa0" not in result
        assert "React Native" in result

    def test_empty_string(self):
        assert _clean_text("") == ""

    def test_preserves_content(self):
        raw = "John Martinez\nSoftware Engineer\n5 years experience"
        result = _clean_text(raw)
        assert "John Martinez" in result
        assert "Software Engineer" in result


# ---------------------------------------------------------------------------
# _detect_format
# ---------------------------------------------------------------------------

class TestDetectFormat:

    def test_pdf_extension(self, tmp_path):
        f = tmp_path / "resume.pdf"
        f.write_bytes(b"%PDF-1.4 fake content")
        assert _detect_format(str(f)) == "pdf"

    def test_docx_extension(self, tmp_path):
        f = tmp_path / "resume.docx"
        f.write_bytes(b"PK\x03\x04 fake docx")
        assert _detect_format(str(f)) == "docx"

    def test_txt_extension(self, tmp_path):
        f = tmp_path / "resume.txt"
        f.write_text("Some resume text")
        assert _detect_format(str(f)) == "txt"

    def test_pdf_magic_bytes_no_extension(self, tmp_path):
        f = tmp_path / "resume"
        f.write_bytes(b"%PDF-1.4 no extension but pdf bytes")
        assert _detect_format(str(f)) == "pdf"

    def test_docx_magic_bytes_no_extension(self, tmp_path):
        f = tmp_path / "myfile"
        f.write_bytes(b"PK\x03\x04 zip-based = docx")
        assert _detect_format(str(f)) == "docx"


# ---------------------------------------------------------------------------
# _split_into_sections
# ---------------------------------------------------------------------------

class TestSplitIntoSections:

    def test_identifies_skills_section(self):
        text = "John Doe\n\nSKILLS\nPython, JavaScript\n\nEXPERIENCE\nEngineer at Company"
        sections = _split_into_sections(text)
        assert "skills" in sections
        assert "Python" in sections["skills"]

    def test_identifies_experience_section(self):
        text = "Jane Smith\n\nEXPERIENCE\nSenior Dev at Corp 2020-2023\n\nEDUCATION\nBS CS 2019"
        sections = _split_into_sections(text)
        assert "experience" in sections
        assert "Senior Dev" in sections["experience"]

    def test_identifies_education_section(self):
        text = "Name\n\nEDUCATION\nBS Computer Science\nState University, 2019"
        sections = _split_into_sections(text)
        assert "education" in sections
        assert "Computer Science" in sections["education"]

    def test_header_captures_pre_section_text(self):
        text = "Alice Johnson\nalice@email.com\n\nSUMMARY\nExperienced engineer"
        sections = _split_into_sections(text)
        assert "header" in sections
        assert "Alice Johnson" in sections["header"]

    def test_handles_colon_after_header(self):
        text = "Name\n\nSkills:\nPython, Java\n\nExperience:\nEngineer"
        sections = _split_into_sections(text)
        assert "skills" in sections

    def test_handles_all_caps_headers(self):
        text = "Name\n\nTECHNICAL SKILLS\nPython\n\nWORK EXPERIENCE\nEngineer"
        sections = _split_into_sections(text)
        assert "skills" in sections
        assert "experience" in sections

    def test_empty_sections_excluded(self):
        text = "Name\n\nSKILLS\n\nEXPERIENCE\nWorked somewhere"
        sections = _split_into_sections(text)
        # Empty skills section should not appear (or be empty string)
        if "skills" in sections:
            assert sections["skills"].strip() != ""

    def test_no_sections_returns_header_only(self):
        text = "Just some random text\nNo section headers here"
        sections = _split_into_sections(text)
        assert "header" in sections


# ---------------------------------------------------------------------------
# extract_resume_text — TXT
# ---------------------------------------------------------------------------

class TestExtractTxt:

    def test_extracts_basic_text(self, tmp_path):
        f = tmp_path / "resume.txt"
        content = "John Doe\n\nSKILLS\nPython, JavaScript\n\nEXPERIENCE\nEngineer at Company for 3 years doing backend work"
        f.write_text(content)
        result = extract_resume_text(str(f))
        assert "John Doe" in result["text"]
        assert result["metadata"]["format"] == "txt"

    def test_returns_sections(self, tmp_path):
        f = tmp_path / "resume.txt"
        f.write_text("Name\n\nSKILLS\nPython Java\n\nEXPERIENCE\nEngineer role performed for two years at a company")
        result = extract_resume_text(str(f))
        assert "sections" in result
        assert isinstance(result["sections"], dict)

    def test_returns_metadata(self, tmp_path):
        f = tmp_path / "resume.txt"
        f.write_text("Name\n\nSKILLS\nPython Java SQL experience databases cloud AWS engineer developer")
        result = extract_resume_text(str(f))
        assert "word_count" in result["metadata"]
        assert "char_count" in result["metadata"]
        assert result["metadata"]["word_count"] > 0

    def test_detects_encoding(self, tmp_path):
        f = tmp_path / "resume.txt"
        # Write latin-1 encoded content
        f.write_bytes("Ren\xe9 Lef\xe8vre\nEngineer in Montr\xe9al working on software applications".encode("latin-1"))
        result = extract_resume_text(str(f))
        assert len(result["text"]) > 0

    def test_raises_on_empty_file(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("")
        with pytest.raises(EmptyDocumentError):
            extract_resume_text(str(f))

    def test_raises_on_too_short_text(self, tmp_path):
        f = tmp_path / "short.txt"
        f.write_text("John Doe")  # only 2 words
        with pytest.raises(EmptyDocumentError):
            extract_resume_text(str(f))

    def test_raises_on_missing_file(self):
        with pytest.raises(FileNotFoundError):
            extract_resume_text("/tmp/definitely_does_not_exist_xyz.txt")

    def test_unicode_bullets_cleaned(self, tmp_path):
        f = tmp_path / "resume.txt"
        f.write_text("Name\n\nSKILLS\n\u2022 Python\n\u2022 JavaScript\n\nEXPERIENCE\nBuilt applications and systems for various clients over multiple years")
        result = extract_resume_text(str(f))
        assert "\u2022" not in result["text"]
        assert "Python" in result["text"]

    def test_full_sample_resume(self):
        sample = os.path.join(os.path.dirname(__file__), "..", "sample_resume.txt")
        if os.path.exists(sample):
            result = extract_resume_text(sample)
            assert result["metadata"]["word_count"] > 100
            assert "skills" in result["sections"]
            assert "experience" in result["sections"]
            assert "education" in result["sections"]


# ---------------------------------------------------------------------------
# extract_resume_text — DOCX
# ---------------------------------------------------------------------------

class TestExtractDocx:

    def _make_docx(self, tmp_path, paragraphs: list[str]) -> str:
        """Helper to create a real .docx file."""
        from docx import Document
        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        path = str(tmp_path / "resume.docx")
        doc.save(path)
        return path

    def test_extracts_paragraphs(self, tmp_path):
        path = self._make_docx(tmp_path, [
            "John Doe", "john@email.com",
            "SKILLS", "Python JavaScript SQL databases cloud AWS engineer developer systems",
            "EXPERIENCE", "Software Engineer at Company for three years building applications"
        ])
        result = extract_resume_text(path)
        assert "John Doe" in result["text"]
        assert result["metadata"]["format"] == "docx"

    def test_sections_from_docx(self, tmp_path):
        path = self._make_docx(tmp_path, [
            "Jane Smith",
            "SKILLS",
            "Python, JavaScript, SQL, React, Node.js working on applications",
            "EXPERIENCE",
            "Senior engineer at TechCorp building scalable systems for enterprise clients"
        ])
        result = extract_resume_text(path)
        assert "skills" in result["sections"] or len(result["sections"]) >= 1

    def test_table_extraction(self, tmp_path):
        """Tables are common in designed resume templates."""
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Python"
        table.cell(0, 1).text = "5 years"
        table.cell(1, 0).text = "JavaScript"
        table.cell(1, 1).text = "4 years"
        # Add enough text to pass word count threshold
        doc.add_paragraph("EXPERIENCE")
        doc.add_paragraph("Software engineer working on web applications and backend services for multiple years at various technology companies building scalable systems")
        path = str(tmp_path / "table_resume.docx")
        doc.save(path)
        result = extract_resume_text(path)
        assert "Python" in result["text"]
        assert "JavaScript" in result["text"]


# ---------------------------------------------------------------------------
# extract_resume_text — PDF
# ---------------------------------------------------------------------------

class TestExtractPdf:

    def _make_simple_pdf(self, tmp_path) -> str:
        """Create a minimal valid PDF with text content."""
        content = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
4 0 obj << /Length 120 >>
stream
BT /F1 12 Tf 72 720 Td (John Doe) Tj 0 -20 Td (SKILLS) Tj 0 -20 Td (Python JavaScript SQL cloud engineer developer systems databases) Tj ET
endstream
endobj
5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000274 00000 n
0000000445 00000 n
trailer << /Size 6 /Root 1 0 R >>
startxref
528
%%EOF"""
        path = str(tmp_path / "resume.pdf")
        with open(path, "wb") as f:
            f.write(content)
        return path

    def test_pdf_format_detected(self, tmp_path):
        path = self._make_simple_pdf(tmp_path)
        assert _detect_format(path) == "pdf"

    def test_pdf_extraction_runs(self, tmp_path):
        """PDF may not yield text in this minimal case — just test no crash."""
        path = self._make_simple_pdf(tmp_path)
        try:
            result = extract_resume_text(path)
            assert "text" in result
            assert "metadata" in result
        except EmptyDocumentError:
            pass  # minimal PDF may not have extractable text — that's fine


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------

class TestEdgeCases:

    def test_oversized_file_raises(self, tmp_path):
        f = tmp_path / "big.txt"
        # Create a file > 10MB
        f.write_bytes(b"x" * (11 * 1024 * 1024))
        with pytest.raises(ResumeExtractionError, match="10MB"):
            extract_resume_text(str(f))

    def test_result_structure_always_complete(self, tmp_path):
        f = tmp_path / "resume.txt"
        f.write_text(
            "Alice Wong\nalice@email.com\n\n"
            "SKILLS\nPython Java SQL cloud AWS engineer developer systems databases\n\n"
            "EXPERIENCE\nSoftware engineer at company for three years building applications"
        )
        result = extract_resume_text(str(f))
        assert "text" in result
        assert "sections" in result
        assert "metadata" in result
        assert isinstance(result["text"], str)
        assert isinstance(result["sections"], dict)
        assert isinstance(result["metadata"], dict)
